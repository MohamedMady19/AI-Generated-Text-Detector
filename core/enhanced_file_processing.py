"""
Enhanced File Processing Module
Handles large files (up to 1 GB) with unlimited processing time
Integrates custom text cleaning methods
"""

import os
import gc
import time
import psutil
import logging
from typing import List, Dict, Tuple, Optional, Generator, Any
from pathlib import Path
import threading
from contextlib import contextmanager

# File processing libraries
import pandas as pd
from docx import Document
import PyPDF2
import chardet

# Custom imports
from .text_cleaning import TextCleaner, split_paragraphs
from ..config import CONFIG, get_memory_limit, should_use_chunked_processing, get_chunk_size

logger = logging.getLogger(__name__)

class MemoryMonitor:
    """Monitor memory usage and trigger cleanup when needed"""
    
    def __init__(self, max_memory_gb: float = 8):
        self.max_memory_bytes = max_memory_gb * 1024 * 1024 * 1024
        self.process = psutil.Process()
    
    def get_memory_usage(self) -> float:
        """Get current memory usage in GB"""
        memory_info = self.process.memory_info()
        return memory_info.rss / (1024 * 1024 * 1024)
    
    def should_cleanup(self) -> bool:
        """Check if memory cleanup is needed"""
        return self.get_memory_usage() > (self.max_memory_bytes / (1024 * 1024 * 1024))
    
    def force_cleanup(self):
        """Force garbage collection and memory cleanup"""
        gc.collect()
        logger.debug(f"Memory cleanup performed. Current usage: {self.get_memory_usage():.2f} GB")

class ProcessingProgress:
    """Track processing progress for large files"""
    
    def __init__(self, total_items: int, log_interval: int = 100):
        self.total_items = total_items
        self.processed_items = 0
        self.log_interval = log_interval
        self.start_time = time.time()
        self.last_log_time = time.time()
        self._lock = threading.Lock()
    
    def update(self, increment: int = 1):
        """Update progress counter"""
        with self._lock:
            self.processed_items += increment
            
            if self.processed_items % self.log_interval == 0:
                self._log_progress()
    
    def _log_progress(self):
        """Log current progress"""
        current_time = time.time()
        elapsed = current_time - self.start_time
        items_per_second = self.processed_items / elapsed if elapsed > 0 else 0
        
        progress_percent = (self.processed_items / self.total_items * 100) if self.total_items > 0 else 0
        
        eta_seconds = (self.total_items - self.processed_items) / items_per_second if items_per_second > 0 else 0
        eta_formatted = f"{eta_seconds/60:.1f} min" if eta_seconds > 60 else f"{eta_seconds:.0f} sec"
        
        logger.info(f"Progress: {self.processed_items}/{self.total_items} ({progress_percent:.1f}%) - "
                   f"{items_per_second:.1f} items/sec - ETA: {eta_formatted}")
    
    def finish(self):
        """Log final completion stats"""
        elapsed = time.time() - self.start_time
        items_per_second = self.processed_items / elapsed if elapsed > 0 else 0
        
        logger.info(f"Processing completed: {self.processed_items} items in {elapsed/60:.1f} minutes "
                   f"({items_per_second:.1f} items/sec)")

class EnhancedFileProcessor:
    """Enhanced file processor for large files with custom text cleaning"""
    
    def __init__(self, config: Dict = None):
        self.config = config or CONFIG
        self.memory_monitor = MemoryMonitor(
            max_memory_gb=self.config.get('MAX_MEMORY_USAGE_GB', 8)
        )
        self.text_cleaner = TextCleaner(
            debug_mode=self.config.get('TEXT_CLEANING_DEBUG_MODE', False),
            save_reports=self.config.get('SAVE_CLEANING_REPORTS', True)
        )
        self.cancelled = False
        self.processing_stats = {
            'files_processed': 0,
            'total_paragraphs': 0,
            'valid_paragraphs': 0,
            'invalid_paragraphs': 0,
            'processing_errors': 0,
            'memory_cleanups': 0
        }
    
    def cancel_processing(self):
        """Cancel ongoing processing"""
        self.cancelled = True
        logger.info("Processing cancellation requested")
    
    def reset_cancellation(self):
        """Reset cancellation flag"""
        self.cancelled = False
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """
        Validate file before processing
        
        Args:
            file_path: Path to file
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return False, "File does not exist"
            
            # Check file size
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            max_size_mb = self.config.get('MAX_FILE_SIZE_MB', 1024)
            
            if file_size_mb > max_size_mb:
                return False, f"File size ({file_size_mb:.1f} MB) exceeds limit ({max_size_mb} MB)"
            
            # Check file extension
            file_ext = Path(file_path).suffix.lower()
            supported_exts = self.config.get('SUPPORTED_EXTENSIONS', ['.txt', '.csv', '.docx', '.pdf'])
            
            if file_ext not in supported_exts:
                return False, f"Unsupported file type: {file_ext}"
            
            # Check if file is readable
            try:
                with open(file_path, 'rb') as f:
                    f.read(1024)  # Try to read first 1KB
            except Exception as e:
                return False, f"File is not readable: {e}"
            
            return True, ""
            
        except Exception as e:
            return False, f"File validation error: {e}"
    
    def detect_encoding(self, file_path: str) -> str:
        """
        Detect file encoding
        
        Args:
            file_path: Path to file
            
        Returns:
            Detected encoding string
        """
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(min(100000, os.path.getsize(file_path)))  # Read up to 100KB for detection
            
            result = chardet.detect(raw_data)
            encoding = result.get('encoding', 'utf-8')
            confidence = result.get('confidence', 0)
            
            logger.debug(f"Detected encoding: {encoding} (confidence: {confidence:.2f})")
            
            # Fallback to common encodings if confidence is low
            if confidence < 0.7:
                for fallback_encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        raw_data.decode(fallback_encoding)
                        logger.debug(f"Using fallback encoding: {fallback_encoding}")
                        return fallback_encoding
                    except UnicodeDecodeError:
                        continue
            
            return encoding or 'utf-8'
            
        except Exception as e:
            logger.warning(f"Encoding detection failed: {e}")
            return 'utf-8'
    
    def read_text_file(self, file_path: str) -> str:
        """
        Read text file with encoding detection and error handling
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content as string
        """
        # Try detected encoding first
        encoding = self.detect_encoding(file_path)
        
        encodings_to_try = [encoding] + [enc for enc in ['utf-8', 'latin-1', 'cp1252', 'utf-16'] 
                                        if enc != encoding]
        
        for enc in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=enc, errors='replace') as f:
                    content = f.read()
                logger.debug(f"Successfully read file with encoding: {enc}")
                return content
            except (UnicodeDecodeError, UnicodeError) as e:
                logger.debug(f"Failed to read with encoding {enc}: {e}")
                continue
        
        raise ValueError(f"Could not read file with any supported encoding: {encodings_to_try}")
    
    def read_pdf_file(self, file_path: str) -> str:
        """
        Read PDF file and extract text
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        try:
            content = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                total_pages = len(pdf_reader.pages)
                logger.info(f"Processing PDF with {total_pages} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    if self.cancelled:
                        break
                    
                    try:
                        page_text = page.extract_text()
                        content += page_text + "\n\n"
                        
                        # Log progress for large PDFs
                        if (page_num + 1) % 50 == 0:
                            logger.info(f"Processed {page_num + 1}/{total_pages} PDF pages")
                        
                        # Memory check for large PDFs
                        if page_num % 100 == 0 and self.memory_monitor.should_cleanup():
                            self.memory_monitor.force_cleanup()
                            self.processing_stats['memory_cleanups'] += 1
                            
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
            
            return content
            
        except Exception as e:
            raise ValueError(f"Failed to read PDF file: {e}")
    
    def read_docx_file(self, file_path: str) -> str:
        """
        Read DOCX file and extract text
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            content = ""
            
            total_paragraphs = len(doc.paragraphs)
            logger.info(f"Processing DOCX with {total_paragraphs} paragraphs")
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                if self.cancelled:
                    break
                
                try:
                    para_text = paragraph.text.strip()
                    if para_text:
                        content += para_text + "\n\n"
                    
                    # Log progress for large documents
                    if (para_num + 1) % 1000 == 0:
                        logger.info(f"Processed {para_num + 1}/{total_paragraphs} DOCX paragraphs")
                    
                    # Memory check
                    if para_num % 500 == 0 and self.memory_monitor.should_cleanup():
                        self.memory_monitor.force_cleanup()
                        self.processing_stats['memory_cleanups'] += 1
                        
                except Exception as e:
                    logger.warning(f"Failed to extract text from paragraph {para_num + 1}: {e}")
                    continue
            
            return content
            
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: {e}")
    
    def read_csv_file(self, file_path: str) -> str:
        """
        Read CSV file and extract text from all columns
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Combined text content from all columns
        """
        try:
            # Try to detect encoding for CSV
            encoding = self.detect_encoding(file_path)
            
            # Read CSV in chunks for large files
            chunk_size = 10000
            content = ""
            
            for chunk_num, chunk in enumerate(pd.read_csv(file_path, encoding=encoding, 
                                                         chunksize=chunk_size, 
                                                         dtype=str, na_filter=False)):
                if self.cancelled:
                    break
                
                # Extract text from all columns
                for col in chunk.columns:
                    col_text = chunk[col].astype(str).str.cat(sep=' ')
                    content += col_text + "\n\n"
                
                logger.info(f"Processed CSV chunk {chunk_num + 1} ({len(chunk)} rows)")
                
                # Memory check
                if chunk_num % 10 == 0 and self.memory_monitor.should_cleanup():
                    self.memory_monitor.force_cleanup()
                    self.processing_stats['memory_cleanups'] += 1
            
            return content
            
        except Exception as e:
            raise ValueError(f"Failed to read CSV file: {e}")
    
    def read_file_content(self, file_path: str) -> str:
        """
        Read file content based on file type
        
        Args:
            file_path: Path to file
            
        Returns:
            File content as string
        """
        file_ext = Path(file_path).suffix.lower()
        
        logger.info(f"Reading file: {file_path} (size: {os.path.getsize(file_path) / (1024*1024):.1f} MB)")
        
        start_time = time.time()
        
        try:
            if file_ext == '.txt':
                content = self.read_text_file(file_path)
            elif file_ext == '.pdf':
                content = self.read_pdf_file(file_path)
            elif file_ext == '.docx':
                content = self.read_docx_file(file_path)
            elif file_ext == '.csv':
                content = self.read_csv_file(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            read_time = time.time() - start_time
            content_size_mb = len(content.encode('utf-8')) / (1024 * 1024)
            
            logger.info(f"File read completed in {read_time:.1f}s. Content size: {content_size_mb:.1f} MB")
            
            return content
            
        except Exception as e:
            logger.error(f"Failed to read file {file_path}: {e}")
            raise
    
    def process_paragraphs_chunked(self, paragraphs: List[str], source_file: str) -> Tuple[List[str], Dict]:
        """
        Process paragraphs in chunks to manage memory usage
        
        Args:
            paragraphs: List of paragraphs to process
            source_file: Source file name for reporting
            
        Returns:
            Tuple of (processed_paragraphs, processing_stats)
        """
        chunk_size = get_chunk_size()
        total_chunks = (len(paragraphs) + chunk_size - 1) // chunk_size
        
        logger.info(f"Processing {len(paragraphs)} paragraphs in {total_chunks} chunks of {chunk_size}")
        
        all_valid_paragraphs = []
        combined_stats = {
            'total_paragraphs': 0,
            'valid_paragraphs': 0,
            'invalid_paragraphs': 0,
            'filter_reasons': {}
        }
        
        progress = ProcessingProgress(total_chunks, log_interval=1)
        
        for chunk_idx in range(total_chunks):
            if self.cancelled:
                logger.info("Processing cancelled by user")
                break
            
            start_idx = chunk_idx * chunk_size
            end_idx = min(start_idx + chunk_size, len(paragraphs))
            chunk_paragraphs = paragraphs[start_idx:end_idx]
            
            logger.debug(f"Processing chunk {chunk_idx + 1}/{total_chunks} "
                        f"(paragraphs {start_idx + 1}-{end_idx})")
            
            try:
                # Process this chunk
                valid_paragraphs, chunk_stats = self.text_cleaner.clean_paragraphs(
                    chunk_paragraphs, 
                    f"{source_file}_chunk_{chunk_idx + 1}"
                )
                
                all_valid_paragraphs.extend(valid_paragraphs)
                
                # Combine stats
                combined_stats['total_paragraphs'] += chunk_stats['total_paragraphs']
                combined_stats['valid_paragraphs'] += chunk_stats['valid_paragraphs']
                combined_stats['invalid_paragraphs'] += chunk_stats['invalid_paragraphs']
                
                for reason, count in chunk_stats['filter_reasons'].items():
                    combined_stats['filter_reasons'][reason] = combined_stats['filter_reasons'].get(reason, 0) + count
                
                progress.update()
                
                # Memory management
                if self.memory_monitor.should_cleanup():
                    self.memory_monitor.force_cleanup()
                    self.processing_stats['memory_cleanups'] += 1
                
                # Give other threads a chance
                time.sleep(0.001)
                
            except Exception as e:
                logger.error(f"Error processing chunk {chunk_idx + 1}: {e}")
                self.processing_stats['processing_errors'] += 1
                continue
        
        progress.finish()
        
        return all_valid_paragraphs, combined_stats
    
    def process_file(self, file_path: str, label: str = "Unknown", source: str = "Unknown") -> Dict[str, Any]:
        """
        Process a single file with enhanced error handling and progress tracking
        
        Args:
            file_path: Path to file
            label: Label for the file (AI Generated, Human Written)
            source: Source of the file (GPT, Claude, Human, etc.)
            
        Returns:
            Dictionary containing processing results
        """
        result = {
            'file_path': file_path,
            'label': label,
            'source': source,
            'success': False,
            'error': None,
            'paragraphs': [],
            'cleaning_stats': {},
            'processing_time': 0,
            'memory_usage_mb': 0
        }
        
        start_time = time.time()
        start_memory = self.memory_monitor.get_memory_usage()
        
        try:
            # Validate file
            is_valid, error_msg = self.validate_file(file_path)
            if not is_valid:
                result['error'] = error_msg
                return result
            
            # Read file content
            content = self.read_file_content(file_path)
            
            if self.cancelled:
                result['error'] = "Processing cancelled"
                return result
            
            # Split into paragraphs
            paragraphs = split_paragraphs(content)
            logger.info(f"Split into {len(paragraphs)} paragraphs")
            
            # Clean paragraphs
            if should_use_chunked_processing() and len(paragraphs) > get_chunk_size():
                valid_paragraphs, cleaning_stats = self.process_paragraphs_chunked(paragraphs, file_path)
            else:
                valid_paragraphs, cleaning_stats = self.text_cleaner.clean_paragraphs(paragraphs, file_path)
            
            result['paragraphs'] = valid_paragraphs
            result['cleaning_stats'] = cleaning_stats
            result['success'] = True
            
            # Update global stats
            self.processing_stats['files_processed'] += 1
            self.processing_stats['total_paragraphs'] += cleaning_stats['total_paragraphs']
            self.processing_stats['valid_paragraphs'] += cleaning_stats['valid_paragraphs']
            self.processing_stats['invalid_paragraphs'] += cleaning_stats['invalid_paragraphs']
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            result['error'] = str(e)
            self.processing_stats['processing_errors'] += 1
        
        finally:
            # Record processing metrics
            result['processing_time'] = time.time() - start_time
            result['memory_usage_mb'] = (self.memory_monitor.get_memory_usage() - start_memory) * 1024
            
            # Final memory cleanup
            self.memory_monitor.force_cleanup()
    
        return result
    
    def get_processing_stats(self) -> Dict:
        """Get current processing statistics"""
        return self.processing_stats.copy()
    
    def reset_stats(self):
        """Reset processing statistics"""
        self.processing_stats = {
            'files_processed': 0,
            'total_paragraphs': 0,
            'valid_paragraphs': 0,
            'invalid_paragraphs': 0,
            'processing_errors': 0,
            'memory_cleanups': 0
        }


@contextmanager
def processing_timeout(seconds: Optional[int]):
    """
    Context manager for processing timeout (disabled if seconds is None)
    
    Args:
        seconds: Timeout in seconds, or None for unlimited
    """
    if seconds is None:
        # No timeout - just yield
        yield
    else:
        # Implement timeout logic here if needed
        # For now, just yield since we want unlimited processing
        yield


def batch_process_files(file_paths: List[str], labels: List[str] = None, 
                       sources: List[str] = None, config: Dict = None) -> Dict[str, Any]:
    """
    Process multiple files in batch with enhanced error handling
    
    Args:
        file_paths: List of file paths to process
        labels: List of labels for each file
        sources: List of sources for each file
        config: Configuration dictionary
        
    Returns:
        Dictionary containing batch processing results
    """
    processor = EnhancedFileProcessor(config)
    
    # Prepare labels and sources
    if labels is None:
        labels = ["Unknown"] * len(file_paths)
    if sources is None:
        sources = ["Unknown"] * len(file_paths)
    
    results = {
        'successful': [],
        'failed': [],
        'processing_stats': {},
        'total_time': 0
    }
    
    start_time = time.time()
    progress = ProcessingProgress(len(file_paths))
    
    try:
        for i, (file_path, label, source) in enumerate(zip(file_paths, labels, sources)):
            if processor.cancelled:
                logger.info("Batch processing cancelled")
                break
            
            logger.info(f"Processing file {i + 1}/{len(file_paths)}: {file_path}")
            
            result = processor.process_file(file_path, label, source)
            
            if result['success']:
                results['successful'].append(result)
            else:
                results['failed'].append(result)
            
            progress.update()
            
            # Memory management between files
            if i % 5 == 0:  # Every 5 files
                processor.memory_monitor.force_cleanup()
    
    finally:
        progress.finish()
        results['total_time'] = time.time() - start_time
        results['processing_stats'] = processor.get_processing_stats()
        
        logger.info(f"Batch processing completed: {len(results['successful'])} successful, "
                   f"{len(results['failed'])} failed in {results['total_time']/60:.1f} minutes")
    
    return results